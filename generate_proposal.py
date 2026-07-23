"""
Generate Proposal Skripsi - Word Document
Sesuai Pedoman Penulisan Tugas Akhir UNM 2019

Jalankan: python generate_proposal.py
Output: Proposal_Skripsi_LMS_AI_Proctoring.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY if isinstance(line_spacing, Pt) else WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing


def set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # Force Times New Roman in East Asian and complex script too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text="", bold=False, italic=False, size=12,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=False,
                  space_before=0, space_after=6, line_rule="double"):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Cm(1.27)
    if line_rule == "double":
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif line_rule == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_rule == "1.5":
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if text:
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return para


def add_heading_bab(doc, text):
    """BAB I, II, III — Times New Roman 14pt Bold Center"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text.upper())
    set_run_font(run, size=14, bold=True)
    return para


def add_heading_sub(doc, text):
    """Sub-bab A. B. C. — Times New Roman 12pt Bold"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text)
    set_run_font(run, size=12, bold=True)
    return para


def add_heading_sub2(doc, text):
    """Sub-sub-bab 1. 2. 3. — Times New Roman 12pt Bold"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text)
    set_run_font(run, size=12, bold=True)
    return para


def add_body(doc, text, first_indent=True, space_before=0):
    """Body text — Times New Roman 12pt, spasi 2"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if first_indent:
        pf.first_line_indent = Cm(1.27)
    run = para.add_run(text)
    set_run_font(run, size=12)
    return para


def add_body_list(doc, items, numbered=True, indent_level=1):
    """Numbered or lettered list body text"""
    for i, item in enumerate(items):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.left_indent = Cm(1.27 * indent_level)
        pf.first_line_indent = Cm(-1.27)
        if numbered:
            prefix = f"{i+1}. "
        else:
            prefix = chr(ord('a') + i) + ". "
        run = para.add_run(prefix + item)
        set_run_font(run, size=12)
    return para


def add_table_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text)
    set_run_font(run, size=12, bold=True)
    return para


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            set_run_font(run, size=11, bold=True)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            for run in row_cells[ci].paragraphs[0].runs:
                set_run_font(run, size=11)
            row_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_separator_line(doc):
    add_paragraph(doc, "─" * 60, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule="single", size=10)


# ─── Main Document Builder ──────────────────────────────────────────────────

def build_proposal():
    doc = Document()

    # ── Page Margins (UNM: atas-kiri 4cm, bawah-kanan 3cm) ──
    for section in doc.sections:
        section.top_margin    = Cm(4)
        section.left_margin   = Cm(4)
        section.bottom_margin = Cm(3)
        section.right_margin  = Cm(3)

    # ── Set default font ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ════════════════════════════════════════════════════════
    # HALAMAN SAMPUL
    # ════════════════════════════════════════════════════════
    p = add_paragraph(doc, "PROPOSAL PENELITIAN",
                      bold=True, size=14,
                      align=WD_ALIGN_PARAGRAPH.CENTER,
                      first_indent=False, line_rule="single",
                      space_before=0, space_after=18)

    add_paragraph(doc,
        "PENGEMBANGAN LEARNING MANAGEMENT SYSTEM DENGAN FITUR AI PROCTORING "
        "DUAL-LAYER BERBASIS YOLOv8 DAN MEDIAPIPE PADA SMA NEGERI 15 MAKASSAR",
        bold=True, size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False, line_rule="single",
        space_before=0, space_after=36)

    add_paragraph(doc, "(Logo UNM — unduh di http://www.unm.ac.id)",
                  italic=True, size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=36)

    add_paragraph(doc, "[NAMA LENGKAP MAHASISWA]",
                  bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=4)

    add_paragraph(doc, "NIM. [NIM MAHASISWA]",
                  bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=36)

    for inst in [
        "Program Studi [Nama Prodi]",
        "Jurusan [Nama Jurusan]",
        "Fakultas [Nama Fakultas]",
        "Universitas Negeri Makassar",
        "[Bulan] [Tahun]"
    ]:
        add_paragraph(doc, inst, size=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER,
                      first_indent=False, line_rule="single",
                      space_before=0, space_after=2)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # HALAMAN PENGESAHAN
    # ════════════════════════════════════════════════════════
    add_paragraph(doc, "PENGESAHAN PROPOSAL",
                  bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=24)

    fields = [
        ("Judul Penelitian",
         "Pengembangan Learning Management System dengan Fitur AI Proctoring "
         "Dual-Layer Berbasis YOLOv8 dan MediaPipe pada SMA Negeri 15 Makassar"),
        ("Nama Mahasiswa", "[Nama Lengkap]"),
        ("NIM", "[NIM]"),
        ("Program Studi", "[Nama Prodi]"),
        ("Jurusan", "[Nama Jurusan]"),
        ("Fakultas", "[Nama Fakultas]"),
    ]
    for label, value in fields:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after = Pt(4)
        run1 = para.add_run(f"{label:<20}: ")
        set_run_font(run1, size=12, bold=True)
        run2 = para.add_run(value)
        set_run_font(run2, size=12)

    add_paragraph(doc, "Makassar, [Tanggal Bulan Tahun]",
                  size=12,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  first_indent=False, line_rule="single",
                  space_before=24, space_after=24)

    # Signature table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    labels_sig = ["Pembimbing I", "Pembimbing II"]
    names_sig  = ["[Nama Pembimbing I]", "[Nama Pembimbing II]"]
    nip_sig    = ["NIP. [NIP]", "NIP. [NIP]"]
    for col in range(2):
        table.cell(0, col).text = labels_sig[col]
        table.cell(1, col).text = ""
        table.cell(2, col).text = ""
        table.cell(3, col).text = names_sig[col]
        for row in range(4):
            for run in table.cell(row, col).paragraphs[0].runs:
                set_run_font(run, size=12, bold=(row in [0, 3]))
            table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "Mengetahui,\nKetua Program Studi [Nama Prodi]",
                  size=12, first_indent=False, line_rule="single",
                  space_before=16, space_after=36)
    add_paragraph(doc, "[Nama Ketua Prodi]\nNIP. [NIP]",
                  size=12, first_indent=False, line_rule="single",
                  space_before=0, space_after=0)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # DAFTAR ISI
    # ════════════════════════════════════════════════════════
    add_paragraph(doc, "DAFTAR ISI",
                  bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=16)

    daftar_isi = [
        ("I. PENDAHULUAN", ""),
        ("    A. Latar Belakang", ""),
        ("    B. Rumusan Masalah", ""),
        ("    C. Tujuan Penelitian", ""),
        ("    D. Manfaat Penelitian", ""),
        ("II. TINJAUAN PUSTAKA", ""),
        ("    A. Kajian Teori", ""),
        ("    B. Kerangka Pikir", ""),
        ("    C. Model Hipotetik", ""),
        ("III. METODE PENELITIAN", ""),
        ("    A. Jenis Penelitian", ""),
        ("    B. Waktu dan Tempat Penelitian", ""),
        ("    C. Desain Penelitian", ""),
        ("    D. Subjek Penelitian", ""),
        ("    E. Definisi Operasional Variabel", ""),
        ("    F. Prosedur Penelitian", ""),
        ("    G. Teknik Pengumpulan Data", ""),
        ("    H. Instrumen Penelitian", ""),
        ("    I. Teknik Analisis Data", ""),
        ("JADWAL RENCANA PELAKSANAAN PENELITIAN", ""),
        ("RENCANA BIAYA PENELITIAN", ""),
        ("DAFTAR PUSTAKA", ""),
    ]
    for item, page in daftar_isi:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        set_run_font(run, size=12)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # BAB I. PENDAHULUAN
    # ════════════════════════════════════════════════════════
    add_heading_bab(doc, "BAB I. PENDAHULUAN")

    add_heading_sub(doc, "A. Latar Belakang")

    latar_belakang = [
        "Perkembangan teknologi informasi dan komunikasi yang pesat telah mendorong "
        "transformasi digital di berbagai sektor kehidupan, termasuk dunia pendidikan. "
        "Era pasca-pandemi COVID-19 mempercepat adopsi pembelajaran berbasis teknologi "
        "di lembaga pendidikan Indonesia, termasuk di tingkat Sekolah Menengah Atas (SMA). "
        "Dalam konteks ini, Learning Management System (LMS) menjadi infrastruktur digital "
        "yang sangat penting untuk mendukung kegiatan belajar mengajar secara terstruktur, "
        "efisien, dan terukur.",

        "SMA Negeri 15 Makassar sebagai salah satu institusi pendidikan menengah di Sulawesi "
        "Selatan menghadapi tantangan dalam mengelola proses pembelajaran dan penilaian secara "
        "digital. Meskipun pelaksanaan ujian berbasis komputer (Computer-Based Test/CBT) telah "
        "mulai diterapkan, praktik kecurangan akademik dalam ujian daring masih menjadi "
        "permasalahan yang belum terselesaikan secara optimal. Kecurangan seperti pertukaran "
        "identitas peserta ujian (joki), penggunaan perangkat komunikasi terlarang seperti ponsel "
        "dan buku, serta ketidakhadiran fisik di depan kamera menjadi tantangan nyata yang "
        "menurunkan validitas dan integritas penilaian.",

        "Berdasarkan hasil observasi awal di SMA Negeri 15 Makassar, diperoleh informasi bahwa "
        "sekolah belum memiliki sistem LMS yang terintegrasi dan belum memiliki mekanisme "
        "pengawasan ujian daring yang dapat diandalkan. Pengawasan ujian masih dilakukan secara "
        "manual oleh pengawas di ruangan, yang tidak dapat memantau seluruh peserta secara "
        "bersamaan. Kondisi ini menciptakan celah yang dapat dimanfaatkan oleh peserta untuk "
        "melakukan kecurangan.",

        "Penelitian terdahulu menunjukkan bahwa sistem proctoring manual memiliki keterbatasan "
        "signifikan dalam mengawasi seluruh peserta ujian secara simultan (Hussain et al., 2019). "
        "Hal ini mendorong pengembangan sistem pengawasan berbasis kecerdasan buatan (Artificial "
        "Intelligence/AI) yang mampu melakukan pemantauan secara otomatis dan real-time. Teknologi "
        "computer vision seperti YOLOv8 (You Only Look Once versi 8) dan MediaPipe telah terbukti "
        "mampu melakukan deteksi objek dan analisis wajah dengan akurasi dan kecepatan tinggi "
        "(Jocher et al., 2023; Lugaresi et al., 2019).",

        "Pada penelitian ini dikembangkan sebuah LMS yang terintegrasi dengan sistem AI Proctoring "
        "berlapis ganda (Dual-Layer). Layer pertama merupakan sistem berbasis server yang menggunakan "
        "YOLOv8 untuk deteksi objek terlarang (ponsel, buku, laptop) serta MediaPipe Face Detection "
        "dan Face Mesh untuk analisis pose kepala dan arah pandangan mata. Layer kedua merupakan "
        "sistem berbasis klien yang berjalan langsung di browser peserta ujian menggunakan face-api.js "
        "dengan model TinyFaceDetector untuk deteksi wajah real-time, verifikasi identitas menggunakan "
        "face descriptor 128 dimensi, serta pendeteksian pola perilaku mencurigakan seperti pola "
        "melihat ponsel.",

        "Berdasarkan uraian di atas, penelitian ini penting untuk dilakukan guna menghasilkan sebuah "
        "LMS yang tidak hanya menyediakan fitur manajemen pembelajaran yang lengkap, tetapi juga "
        "dilengkapi dengan sistem pengawasan ujian berbasis AI yang andal dan terverifikasi "
        "kelayakannya. Oleh karena itu, penelitian ini mengangkat judul \u201cPengembangan Learning "
        "Management System dengan Fitur AI Proctoring Dual-Layer Berbasis YOLOv8 dan MediaPipe "
        "pada SMA Negeri 15 Makassar\u201d.",
    ]
    for para_text in latar_belakang:
        add_body(doc, para_text, first_indent=True)

    add_heading_sub(doc, "B. Rumusan Masalah")
    add_body(doc, "Berdasarkan uraian latar belakang di atas, rumusan masalah dalam penelitian ini adalah:", first_indent=True)
    rumusan = [
        "Bagaimana merancang dan mengembangkan Learning Management System yang mencakup fitur "
        "manajemen pembelajaran, absensi QR dinamis, dan ujian CBT untuk SMA Negeri 15 Makassar?",
        "Bagaimana mengimplementasikan sistem AI Proctoring Dual-Layer berbasis YOLOv8 dan "
        "MediaPipe yang mampu mendeteksi kecurangan ujian secara otomatis dan real-time?",
        "Bagaimana tingkat kelayakan Learning Management System dengan fitur AI Proctoring "
        "Dual-Layer yang dikembangkan berdasarkan penilaian ahli media, ahli materi, dan "
        "respon pengguna?",
    ]
    add_body_list(doc, rumusan, numbered=True)

    add_heading_sub(doc, "C. Tujuan Penelitian")
    add_body(doc, "Berdasarkan rumusan masalah di atas, tujuan dari penelitian ini adalah:", first_indent=True)
    tujuan = [
        "Merancang dan mengembangkan Learning Management System yang terintegrasi dengan fitur "
        "manajemen pembelajaran, absensi QR dinamis, dan ujian CBT di SMA Negeri 15 Makassar.",
        "Mengimplementasikan sistem AI Proctoring Dual-Layer berbasis YOLOv8 dan MediaPipe yang "
        "dapat mendeteksi kecurangan ujian secara otomatis dalam mode real-time.",
        "Menguji dan mengevaluasi tingkat kelayakan sistem yang dikembangkan berdasarkan aspek "
        "fungsionalitas, kegunaan (usability), dan keakuratan deteksi proctoring.",
    ]
    add_body_list(doc, tujuan, numbered=True)

    add_heading_sub(doc, "D. Manfaat Penelitian")
    add_heading_sub2(doc, "1. Manfaat Teoritis")
    manfaat_teoritis = [
        "Memberikan kontribusi ilmiah pada pengembangan sistem manajemen pembelajaran berbasis AI, "
        "khususnya pada aspek pengawasan ujian daring menggunakan arsitektur dual-layer.",
        "Menjadi referensi bagi penelitian selanjutnya mengenai implementasi YOLOv8 dan MediaPipe "
        "dalam sistem pendidikan berbasis web.",
        "Menambah khasanah pengetahuan tentang arsitektur sistem proctoring yang mengintegrasikan "
        "deteksi berbasis server dan berbasis klien secara sinergis.",
    ]
    add_body_list(doc, manfaat_teoritis, numbered=False)

    add_heading_sub2(doc, "2. Manfaat Praktis")
    manfaat_praktis = [
        "Bagi Sekolah: tersedianya sistem LMS yang modern dan terintegrasi untuk mendukung proses "
        "pembelajaran digital, absensi, dan evaluasi di SMA Negeri 15 Makassar.",
        "Bagi Guru: memudahkan pembuatan soal ujian, pengelolaan sesi ujian, dan pemantauan "
        "integritas ujian peserta didik secara real-time melalui dashboard monitoring proctoring.",
        "Bagi Siswa: memberikan pengalaman ujian digital yang terstruktur, adil, dan setara "
        "bagi seluruh peserta, dengan antarmuka yang intuitif dan responsif.",
        "Bagi Peneliti dan Pengembang: menjadi referensi implementasi arsitektur AI proctoring "
        "dual-layer pada sistem pendidikan, termasuk konfigurasi threshold deteksi yang optimal.",
    ]
    add_body_list(doc, manfaat_praktis, numbered=False)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # BAB II. TINJAUAN PUSTAKA
    # ════════════════════════════════════════════════════════
    add_heading_bab(doc, "BAB II. TINJAUAN PUSTAKA")
    add_heading_sub(doc, "A. Kajian Teori")

    teori_items = [
        ("1. Learning Management System (LMS)",
         "Learning Management System (LMS) adalah platform perangkat lunak yang dirancang untuk "
         "membuat, mendistribusikan, mengelola, dan menilai konten pembelajaran secara digital. "
         "Ellis (2009) mendefinisikan LMS sebagai infrastruktur perangkat lunak yang mengelola "
         "konten dan sumber daya pembelajaran secara online, yang mencakup manajemen pengguna, "
         "manajemen konten, penilaian, komunikasi, dan pelacakan kemajuan belajar. Dalam penelitian "
         "ini, LMS yang dikembangkan mencakup tiga modul utama: (1) manajemen pembelajaran yang "
         "meliputi distribusi materi, tugas, pengumuman, dan rekap nilai; (2) absensi QR dinamis "
         "dengan mekanisme anti-titip absen; dan (3) ujian CBT yang terintegrasi dengan sistem AI Proctoring."),

        ("2. Computer-Based Test (CBT)",
         "Computer-Based Test (CBT) adalah metode pelaksanaan ujian yang menggunakan komputer "
         "atau perangkat digital sebagai media utama ujian (Mardapi, 2017). CBT menawarkan "
         "keunggulan dibandingkan ujian berbasis kertas antara lain efisiensi distribusi soal, "
         "penilaian otomatis, pengacakan soal dan pilihan jawaban, penghematan kertas, serta "
         "kemudahan penyimpanan dan pengolahan data hasil ujian. Namun, CBT juga menghadirkan "
         "tantangan terkait integritas ujian. Tanpa pengawasan yang memadai, peserta dapat "
         "memanfaatkan celah dalam sistem untuk melakukan kecurangan seperti membuka tab browser "
         "lain, menggunakan perangkat komunikasi, atau menukar identitas (Raman et al., 2019)."),

        ("3. AI Proctoring",
         "AI Proctoring adalah pemanfaatan kecerdasan buatan untuk mengotomatiskan proses "
         "pengawasan ujian online (Hussain et al., 2019). Sistem ini memanfaatkan teknologi "
         "computer vision, machine learning, dan analisis perilaku untuk mendeteksi aktivitas "
         "mencurigakan selama ujian berlangsung tanpa memerlukan kehadiran pengawas manusia "
         "secara penuh. Jenis deteksi yang diterapkan meliputi: (a) deteksi kehadiran wajah "
         "(face presence detection); (b) deteksi wajah ganda (multiple face detection); "
         "(c) analisis pose kepala (head pose estimation); (d) analisis arah pandangan mata "
         "(eye gaze estimation); (e) verifikasi identitas (identity verification); dan (f) "
         "deteksi objek terlarang (prohibited object detection)."),

        ("4. YOLOv8 (You Only Look Once Versi 8)",
         "YOLOv8 adalah generasi terbaru dari arsitektur deteksi objek real-time YOLO yang "
         "dikembangkan oleh Ultralytics (Jocher et al., 2023). YOLOv8 mengadopsi arsitektur "
         "anchor-free dengan backbone CSP (Cross-Stage Partial) yang dioptimalkan untuk "
         "keseimbangan antara kecepatan inferensi dan akurasi deteksi. Dalam penelitian ini "
         "digunakan model YOLOv8 nano (yolov8n.pt) yang dilatih pada dataset COCO dengan 80 "
         "kelas objek. Kelas objek yang dideteksi sebagai terlarang dalam konteks ujian "
         "adalah: cell phone (ID 67), book (ID 73), laptop (ID 63), dan tv/monitor (ID 62). "
         "Menurut Terven et al. (2023), YOLOv8 nano mencapai nilai mAP sebesar 37,3% pada "
         "dataset COCO val2017 dengan kecepatan inferensi 80,4 FPS pada GPU."),

        ("5. MediaPipe",
         "MediaPipe adalah framework machine learning cross-platform yang dikembangkan oleh "
         "Google untuk membangun pipeline persepsi multimodal pada perangkat edge dan server "
         "(Lugaresi et al., 2019). Dalam penelitian ini digunakan dua solusi MediaPipe: "
         "(a) MediaPipe Face Detection menggunakan model BlazeFace berbasis MobileNetV2 untuk "
         "mendeteksi wajah dengan confidence tinggi pada jarak pendek; dan (b) MediaPipe Face "
         "Mesh yang menghasilkan 468 titik landmark wajah 3D termasuk iris landmarks (468-477) "
         "untuk estimasi pose kepala menggunakan algoritma cv2.solvePnP dengan threshold yaw "
         "38 derajat dan pitch 33 derajat, serta estimasi arah pandangan mata menggunakan "
         "posisi relatif iris dengan threshold deviasi 0,48. Sistem juga mengimplementasikan "
         "mekanisme image enhancement berlapis (CLAHE, Gamma Correction, dan kombinasi keduanya) "
         "sebagai fallback ketika wajah tidak terdeteksi pada gambar asli."),

        ("6. face-api.js",
         "face-api.js adalah pustaka JavaScript yang mengimplementasikan beberapa arsitektur "
         "neural network untuk deteksi dan analisis wajah langsung di browser menggunakan "
         "TensorFlow.js sebagai backend komputasi (Jurkiewicz, 2018). Model yang digunakan "
         "dalam penelitian ini adalah: (a) TinyFaceDetector, model deteksi wajah ringan berbasis "
         "MobileNetV1 dengan input size 224px (mobile) atau 320px (desktop) dan score threshold "
         "0,35; (b) FaceLandmark68TinyNet, model yang menghasilkan 68 titik facial landmark untuk "
         "analisis geometri wajah; dan (c) FaceRecognitionNet berbasis ResNet-34 yang menghasilkan "
         "face descriptor 128 dimensi untuk verifikasi identitas menggunakan euclidean distance "
         "dengan threshold 0,5."),

        ("7. Research and Development (R&D)",
         "Research and Development (R&D) adalah metode penelitian yang digunakan untuk menghasilkan "
         "produk tertentu dan menguji efektivitas produk tersebut (Sugiyono, 2019). Model R&D "
         "yang digunakan mengacu pada model Borg & Gall (1983) yang diadaptasi menjadi tiga "
         "tahap utama yaitu Define (pendefinisian), Design (perancangan), dan Develop "
         "(pengembangan), yang mengacu pada model 4-D (Four-D) yang dikembangkan oleh "
         "Thiagarajan et al. (1974) dan diadaptasi untuk pengembangan sistem perangkat lunak."),
    ]

    for title, content in teori_items:
        add_heading_sub2(doc, title)
        add_body(doc, content, first_indent=True)

    # Tabel Dual-Layer
    add_heading_sub2(doc, "8. Arsitektur Dual-Layer Proctoring")
    add_body(doc,
        "Kebaruan (novelty) utama penelitian ini terletak pada implementasi arsitektur proctoring "
        "berlapis ganda yang bekerja secara sinergis dan saling melengkapi. Layer pertama berjalan "
        "di server dan dieksekusi setiap kali snapshot berkala dikirim (interval 10-15 detik), "
        "sementara layer kedua berjalan langsung di browser peserta ujian dengan interval deteksi "
        "1,5 detik untuk pemantauan real-time yang lebih responsif.",
        first_indent=True)

    add_table_caption(doc, "Tabel 2.1 Perbandingan Karakteristik Dual-Layer AI Proctoring")
    add_simple_table(doc,
        headers=["Aspek", "Layer 1 (Server-Side)", "Layer 2 (Client-Side)"],
        rows=[
            ["Teknologi", "YOLOv8 + MediaPipe (Python)", "face-api.js (JavaScript)"],
            ["Eksekusi", "Server FastAPI", "Browser Peserta"],
            ["Interval Analisis", "10-15 detik", "1,5 detik"],
            ["Deteksi Objek Terlarang", "YOLOv8 (80 kelas COCO)", "Tidak tersedia"],
            ["Deteksi Wajah", "MediaPipe BlazeFace", "TinyFaceDetector"],
            ["Pose Kepala", "solvePnP (sudut Euler)", "Geometri landmark 68 titik"],
            ["Eye Gaze", "Iris landmark MediaPipe", "Rasio posisi pupil"],
            ["Verifikasi Identitas", "dlib ResNet 128-dim", "FaceRecognitionNet 128-dim"],
            ["Deteksi Pola HP", "Tidak tersedia", "Absence cycle detection"],
            ["Konfirmasi Temporal", "Dedup window 15 detik", "Window 10 detik + cooldown"],
        ]
    )
    doc.add_paragraph()  # spacing after table

    add_heading_sub(doc, "B. Kerangka Pikir")
    add_body(doc,
        "SMA Negeri 15 Makassar menghadapi dua permasalahan utama: (1) belum tersedianya sistem "
        "manajemen pembelajaran digital yang terintegrasi, dan (2) belum adanya mekanisme "
        "pengawasan ujian berbasis AI yang dapat memastikan integritas ujian daring. Kondisi ini "
        "menyebabkan proses pembelajaran dan evaluasi berjalan kurang efisien dan rentan terhadap "
        "kecurangan akademik.",
        first_indent=True)
    add_body(doc,
        "Untuk mengatasi kedua permasalahan tersebut, dikembangkan sebuah LMS terintegrasi dengan "
        "sistem AI Proctoring Dual-Layer. Sistem AI Proctoring Dual-Layer memanfaatkan YOLOv8 dan "
        "MediaPipe pada sisi server untuk analisis mendalam setiap snapshot berkala, serta "
        "face-api.js pada sisi klien untuk pemantauan real-time di browser peserta. Kedua layer "
        "bekerja secara sinergis: layer klien memberikan umpan balik cepat (1,5 detik) untuk "
        "pelanggaran berbasis perilaku wajah, sementara layer server memberikan deteksi komprehensif "
        "termasuk objek terlarang dan verifikasi identitas berbasis face embedding. Hasil deteksi "
        "dari kedua layer disimpan di database, dikumulasikan menjadi risk score komposit (0-100), "
        "dan ditampilkan secara real-time pada dashboard monitoring guru melalui Socket.io.",
        first_indent=True)

    add_heading_sub(doc, "C. Model Hipotetik")
    add_body(doc,
        "Berdasarkan kajian teori dan kerangka pikir di atas, model hipotetik sistem yang "
        "dikembangkan menggambarkan alur kerja lengkap dari browser siswa hingga dashboard guru. "
        "Layer kedua (face-api.js) berjalan di browser siswa setiap 1,5 detik mendeteksi "
        "no_face, multi_face, head_turn, eye_gaze, identity_mismatch, dan suspect_phone_check, "
        "kemudian mengirim laporan pelanggaran ke Laravel Backend via REST API. Layer pertama "
        "(YOLOv8 + MediaPipe) berjalan di FastAPI microservice setiap snapshot dikirim, "
        "menghitung risk score komposit, dan hasilnya di-broadcast ke dashboard monitoring guru "
        "melalui Socket.io secara real-time.",
        first_indent=True)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # BAB III. METODE PENELITIAN
    # ════════════════════════════════════════════════════════
    add_heading_bab(doc, "BAB III. METODE PENELITIAN")

    add_heading_sub(doc, "A. Jenis Penelitian")
    add_body(doc,
        "Penelitian ini merupakan penelitian pengembangan dengan menggunakan metode Research and "
        "Development (R&D). R&D dipilih karena tujuan utama penelitian adalah menghasilkan produk "
        "berupa Learning Management System dengan fitur AI Proctoring Dual-Layer yang telah diuji "
        "tingkat kelayakannya. Model pengembangan yang digunakan mengacu pada model Borg & Gall "
        "yang diadaptasi menjadi tiga tahap utama, yaitu: (1) Pendefinisian (Define), "
        "(2) Perancangan (Design), dan (3) Pengembangan (Develop).",
        first_indent=True)

    add_heading_sub(doc, "B. Waktu dan Tempat Penelitian")
    add_heading_sub2(doc, "1. Waktu Penelitian")
    add_body(doc,
        "Penelitian ini direncanakan dilaksanakan selama [X] bulan, mulai dari bulan [Bulan Mulai] "
        "sampai bulan [Bulan Selesai] tahun [Tahun].",
        first_indent=True)
    add_heading_sub2(doc, "2. Tempat Penelitian")
    add_body(doc,
        "Penelitian ini dilaksanakan di SMA Negeri 15 Makassar yang beralamat di [Alamat Lengkap "
        "SMA Negeri 15 Makassar]. Pengembangan sistem dilakukan di [Tempat Pengembangan, misalnya: "
        "laboratorium komputer/tempat tinggal peneliti].",
        first_indent=True)

    add_heading_sub(doc, "C. Desain Penelitian")
    add_body(doc,
        "Desain penelitian yang digunakan mengacu pada model pengembangan Borg & Gall yang "
        "diadaptasi menjadi tiga tahap utama. Tahap pertama adalah Pendefinisian (Define) yang "
        "meliputi analisis kebutuhan, analisis pengguna, dan analisis tugas. Tahap kedua adalah "
        "Perancangan (Design) yang meliputi perancangan arsitektur sistem, database, antarmuka, "
        "dan alur AI Proctoring dual-layer. Tahap ketiga adalah Pengembangan (Develop) yang "
        "meliputi implementasi kode, pengujian fungsional, validasi ahli, revisi produk, dan "
        "uji coba pengguna.",
        first_indent=True)

    add_heading_sub(doc, "D. Subjek Penelitian")
    add_heading_sub2(doc, "1. Subjek Validasi Ahli")
    add_body_list(doc, [
        "Ahli Media: dosen atau praktisi yang memiliki kompetensi di bidang teknologi informasi "
        "dan pengembangan sistem berbasis web, berjumlah [X] orang.",
        "Ahli Materi: dosen atau guru yang memiliki kompetensi di bidang teknologi pembelajaran "
        "dan sistem evaluasi pendidikan, berjumlah [X] orang.",
    ], numbered=False)
    add_heading_sub2(doc, "2. Subjek Uji Coba Pengguna")
    add_body_list(doc, [
        "Guru: guru di SMA Negeri 15 Makassar yang terlibat dalam proses pembelajaran dan "
        "pelaksanaan ujian, berjumlah [X] orang.",
        "Siswa: siswa SMA Negeri 15 Makassar yang dipilih secara purposive sampling, "
        "berjumlah [X] orang.",
    ], numbered=False)

    add_heading_sub(doc, "E. Definisi Operasional Variabel")
    def_items = [
        ("1. Learning Management System (LMS)",
         "Dalam penelitian ini, LMS adalah sistem perangkat lunak berbasis web yang dikembangkan "
         "menggunakan Next.js 14 (frontend) dan Laravel 11 (backend) yang menyediakan fitur: "
         "(a) manajemen pengguna multi-peran (Admin, Guru, Siswa); (b) absensi QR dinamis dengan "
         "QR Code yang diperbarui setiap 2-5 menit; (c) distribusi materi pembelajaran; "
         "(d) pengelolaan tugas; (e) ujian CBT dengan mekanisme anti-cheat; dan "
         "(f) pelaporan nilai dan kehadiran."),
        ("2. AI Proctoring Dual-Layer",
         "AI Proctoring Dual-Layer dalam penelitian ini adalah sistem pengawasan ujian berbasis "
         "kecerdasan buatan yang terdiri dari dua lapisan: (a) Layer Server-Side yang menggunakan "
         "YOLOv8 nano untuk deteksi objek terlarang dan MediaPipe untuk analisis pose kepala serta "
         "arah pandangan mata, dieksekusi setiap 10-15 detik; (b) Layer Client-Side yang menggunakan "
         "face-api.js untuk deteksi wajah real-time, verifikasi identitas, dan deteksi pola perilaku "
         "mencurigakan, dieksekusi setiap 1,5 detik."),
        ("3. Kelayakan Sistem",
         "Kelayakan sistem dalam penelitian ini merujuk pada tingkat kesesuaian sistem yang "
         "dikembangkan dengan kebutuhan pengguna dan standar kualitas perangkat lunak, diukur "
         "melalui: (a) validasi ahli media dan ahli materi menggunakan angket berskala Likert; "
         "(b) respon pengguna (guru dan siswa) menggunakan angket berskala Likert; dan (c) akurasi "
         "sistem AI Proctoring yang diukur menggunakan metrik precision, recall, dan F1-score."),
    ]
    for title, content in def_items:
        add_heading_sub2(doc, title)
        add_body(doc, content, first_indent=True)

    add_heading_sub(doc, "F. Prosedur Penelitian")
    prosedur_items = [
        ("1. Tahap Pendefinisian (Define)",
         ["Analisis Kebutuhan (Needs Analysis): dilakukan melalui observasi langsung di SMA Negeri "
          "15 Makassar dan wawancara dengan guru, staf TI, dan kepala sekolah untuk mengidentifikasi "
          "permasalahan dalam pengelolaan pembelajaran dan pelaksanaan ujian digital.",
          "Analisis Pengguna (User Analysis): identifikasi karakteristik dan kebutuhan tiga peran "
          "pengguna yaitu Admin (mengelola pengguna, kelas, jadwal, statistik), Guru (membuat sesi "
          "absensi, mengelola bank soal, membuat dan memantau ujian), dan Siswa (melakukan absensi, "
          "mengerjakan ujian, mengakses materi dan nilai).",
          "Analisis Tugas (Task Analysis): pemetaan alur proses bisnis utama meliputi alur absensi "
          "QR dinamis, alur pelaksanaan ujian CBT dengan anti-cheat, dan alur AI Proctoring dual-layer."]),
        ("2. Tahap Perancangan (Design)",
         ["Perancangan Arsitektur Sistem: sistem dibangun dengan arsitektur microservices yang "
          "terdiri dari Next.js 14 (frontend, port 3000), Laravel 11 (backend API, port 8000), "
          "FastAPI Python (microservice AI, port 8001), Socket.io server (port 6001), dan "
          "PostgreSQL (Supabase) sebagai database.",
          "Perancangan Database: merancang skema relasional yang mencakup tabel users, classes, "
          "subjects, exams, questions, question_banks, exam_results (dengan field baseline_face_embedding "
          "JSON 128 dimensi), proctoring_alerts, proctoring_scores, monitoring_snapshots, "
          "attendance_sessions, dan attendances.",
          "Perancangan Antarmuka (UI/UX): merancang antarmuka responsif menggunakan Tailwind CSS "
          "yang mencakup dashboard multi-role, halaman ujian CBT dengan fullscreen lock, panel "
          "monitoring proctoring real-time, dan halaman diagnostik proctoring.",
          "Perancangan Alur AI Proctoring Dual-Layer: merancang alur lengkap dari upload snapshot "
          "hingga broadcast alert ke dashboard guru, termasuk mekanisme konfirmasi temporal untuk "
          "mengurangi false positive."]),
        ("3. Tahap Pengembangan (Develop)",
         ["Implementasi Sistem: mengimplementasikan seluruh komponen sistem berdasarkan desain "
          "yang telah dirancang menggunakan stack teknologi yang telah ditentukan.",
          "Pengujian Fungsional: menguji setiap fitur sistem untuk memastikan berfungsi sesuai "
          "dengan spesifikasi yang telah dirancang.",
          "Validasi Ahli: instrumen validasi diberikan kepada ahli media dan ahli materi untuk "
          "mendapatkan penilaian kelayakan sistem dan masukan perbaikan.",
          "Revisi Produk: melakukan perbaikan sistem berdasarkan masukan dari validator.",
          "Uji Coba Pengguna: sistem yang telah direvisi diujicobakan kepada guru dan siswa.",
          "Produk Akhir: menghasilkan sistem LMS dengan AI Proctoring Dual-Layer yang telah "
          "divalidasi dan siap digunakan di SMA Negeri 15 Makassar."]),
    ]
    for title, items in prosedur_items:
        add_heading_sub2(doc, title)
        add_body_list(doc, items, numbered=False)

    add_heading_sub(doc, "G. Teknik Pengumpulan Data")
    add_body_list(doc, [
        "Angket (Kuesioner): digunakan untuk mengumpulkan data kelayakan sistem dari ahli media, "
        "ahli materi, dan pengguna. Angket menggunakan skala Likert dengan 5 alternatif jawaban: "
        "Sangat Baik (5), Baik (4), Cukup (3), Kurang (2), dan Sangat Kurang (1).",
        "Observasi: dilakukan untuk mengidentifikasi kebutuhan sistem pada tahap pendefinisian "
        "dan mencatat data deteksi AI Proctoring (TP, FP, TN, FN) pada tahap pengujian akurasi.",
        "Wawancara: dilakukan kepada guru dan kepala sekolah pada tahap analisis kebutuhan untuk "
        "menggali informasi tentang kondisi pembelajaran di SMA Negeri 15 Makassar.",
        "Dokumentasi: digunakan untuk mengumpulkan data berupa tangkapan layar antarmuka sistem, "
        "log pelanggaran proctoring, dan foto pelaksanaan uji coba.",
    ], numbered=True)

    add_heading_sub(doc, "H. Instrumen Penelitian")
    add_body_list(doc, [
        "Lembar Validasi Ahli Media: mengevaluasi aspek desain antarmuka, kemudahan navigasi "
        "(usability), keterbacaan, responsivitas pada berbagai perangkat, dan performa sistem.",
        "Lembar Validasi Ahli Materi: mengevaluasi kesesuaian fitur LMS dengan kebutuhan "
        "pembelajaran, kesesuaian mekanisme ujian CBT, dan ketepatan fitur proctoring.",
        "Angket Respon Pengguna: mengevaluasi kemudahan penggunaan, manfaat yang dirasakan, "
        "kepuasan terhadap fitur, dan kesesuaian sistem dengan kebutuhan nyata.",
        "Lembar Observasi Akurasi Proctoring: mencatat hasil deteksi dalam bentuk matriks konfusi "
        "(TP, FP, TN, FN) untuk setiap jenis deteksi AI.",
    ], numbered=True)

    add_heading_sub(doc, "I. Teknik Analisis Data")
    add_heading_sub2(doc, "1. Analisis Kelayakan Sistem")
    add_body(doc,
        "Data kuantitatif dari angket validasi ahli dan respon pengguna dianalisis menggunakan "
        "perhitungan persentase kelayakan dengan rumus: Persentase = (Skor Aktual / Skor Maksimal) "
        "\u00d7 100%. Hasil persentase kemudian diinterpretasikan berdasarkan kriteria kelayakan "
        "menurut Riduwan (2015). Sistem dinyatakan layak digunakan apabila rata-rata persentase "
        "kelayakan dari seluruh aspek mencapai minimal 61% (kategori Layak).",
        first_indent=True)

    add_table_caption(doc, "Tabel 3.1 Kriteria Kelayakan Sistem")
    add_simple_table(doc,
        headers=["Persentase", "Kriteria"],
        rows=[
            ["81% - 100%", "Sangat Layak"],
            ["61% - 80%",  "Layak"],
            ["41% - 60%",  "Cukup Layak"],
            ["21% - 40%",  "Kurang Layak"],
            ["0% - 20%",   "Tidak Layak"],
        ]
    )
    doc.add_paragraph()

    add_heading_sub2(doc, "2. Analisis Akurasi Sistem AI Proctoring")
    add_body(doc,
        "Akurasi sistem AI Proctoring dianalisis menggunakan metrik evaluasi klasifikasi biner "
        "berdasarkan matriks konfusi, meliputi: (a) Accuracy = (TP + TN) / (TP + TN + FP + FN); "
        "(b) Precision = TP / (TP + FP); (c) Recall = TP / (TP + FN); dan "
        "(d) F1-Score = 2 \u00d7 (Precision \u00d7 Recall) / (Precision + Recall). "
        "Keterangan: TP (True Positive) = kecurangan terdeteksi oleh sistem dan memang terjadi; "
        "TN (True Negative) = tidak ada kecurangan dan sistem tidak mendeteksi; "
        "FP (False Positive) = sistem mendeteksi kecurangan padahal tidak terjadi; "
        "FN (False Negative) = kecurangan terjadi namun tidak terdeteksi sistem.",
        first_indent=True)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # JADWAL PENELITIAN
    # ════════════════════════════════════════════════════════
    add_paragraph(doc, "JADWAL RENCANA PELAKSANAAN PENELITIAN",
                  bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=16)

    add_table_caption(doc, "Tabel Jadwal Rencana Pelaksanaan Penelitian")
    add_simple_table(doc,
        headers=["No", "Kegiatan", "Bln 1", "Bln 2", "Bln 3", "Bln 4", "Bln 5", "Bln 6"],
        rows=[
            ["1", "Studi Literatur & Analisis Kebutuhan", "\u2713", "", "", "", "", ""],
            ["2", "Perancangan Sistem (Design)", "\u2713", "\u2713", "", "", "", ""],
            ["3", "Implementasi Sistem (Develop)", "", "\u2713", "\u2713", "\u2713", "", ""],
            ["4", "Pengujian Fungsional", "", "", "", "\u2713", "", ""],
            ["5", "Validasi Ahli Media dan Materi", "", "", "", "\u2713", "", ""],
            ["6", "Revisi Produk", "", "", "", "", "\u2713", ""],
            ["7", "Uji Coba Pengguna", "", "", "", "", "\u2713", ""],
            ["8", "Analisis Data dan Pembahasan", "", "", "", "", "\u2713", "\u2713"],
            ["9", "Penulisan Laporan Skripsi", "", "", "", "", "\u2713", "\u2713"],
            ["10","Seminar Hasil dan Ujian Tutup", "", "", "", "", "", "\u2713"],
        ]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # RENCANA BIAYA
    # ════════════════════════════════════════════════════════
    add_paragraph(doc, "RENCANA BIAYA PENELITIAN",
                  bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=16)

    add_table_caption(doc, "Tabel Rencana Biaya Penelitian")
    add_simple_table(doc,
        headers=["No", "Uraian Kegiatan", "Satuan", "Jumlah", "Harga Satuan (Rp)", "Total (Rp)"],
        rows=[
            ["", "A. Biaya Operasional", "", "", "", ""],
            ["1", "Sewa VPS/Hosting (6 bulan)", "Bulan", "6", "...", "..."],
            ["2", "Transportasi ke sekolah", "Kali", "...", "...", "..."],
            ["3", "Koneksi internet", "Bulan", "6", "...", "..."],
            ["", "B. ATK & Percetakan", "", "", "", ""],
            ["4", "Cetak instrumen penelitian", "Eksemplar", "...", "...", "..."],
            ["5", "Cetak draft dan laporan skripsi", "Eksemplar", "...", "...", "..."],
            ["6", "Penjilidan laporan", "Eksemplar", "...", "...", "..."],
            ["", "C. Biaya Tak Terduga", "", "", "", "..."],
            ["", "TOTAL", "", "", "", "..."],
        ]
    )
    add_body(doc, "Sumber Dana: Biaya Mandiri Peneliti", first_indent=False)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    # DAFTAR PUSTAKA
    # ════════════════════════════════════════════════════════
    add_paragraph(doc, "DAFTAR PUSTAKA",
                  bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False, line_rule="single",
                  space_before=0, space_after=16)

    references = [
        "Borg, W. R., & Gall, M. D. (1983). Educational research: An introduction (4th ed.). Longman.",
        "Ellis, R. K. (2009). A field guide to learning management systems. American Society for Training and Development (ASTD).",
        "Hussain, M., Babe, A., Sarawgi, S., Rathod, V., Bhatt, N., & Gharpure, P. (2019). Automated proctoring systems: A review. International Journal of Scientific Research in Science and Technology, 6(2), 314-321.",
        "Jocher, G., Chaurasia, A., & Qiu, J. (2023). Ultralytics YOLOv8. https://github.com/ultralytics/ultralytics",
        "Jurkiewicz, V. (2018). face-api.js \u2014 JavaScript API for face recognition in the browser with tensorflow.js. https://github.com/justadudewhohacks/face-api.js",
        "Lugaresi, C., Tang, J., Nash, H., McClanahan, C., Uboweja, E., Hays, M., Zhang, F., Chang, C. L., Yong, M. G., Lee, J., Chang, W. T., Hua, W., Georg, M., & Grundmann, M. (2019). MediaPipe: A framework for perceiving and processing reality. Third Workshop on Computer Vision for AR/VR at IEEE CVPR 2019.",
        "Mardapi, D. (2017). Pengukuran, penilaian, dan evaluasi pendidikan. Parama Publishing.",
        "Moodle. (2020). About Moodle: Learning management system. https://moodle.org/about/",
        "Raman, R., Vachharajani, H., & Nedungadi, P. (2019). Adoption of online proctored examinations by university students during COVID-19: Innovation diffusion study. Education and Information Technologies, 26(1), 5865-5876. https://doi.org/10.1007/s10639-021-10581-5",
        "Redmon, J., Divvala, S., Girshick, R., & Farhadi, A. (2016). You only look once: Unified, real-time object detection. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 779-788. https://doi.org/10.1109/CVPR.2016.91",
        "Riduwan. (2015). Dasar-dasar statistika. Alfabeta.",
        "Sugiyono. (2019). Metode penelitian dan pengembangan (Research and Development). Alfabeta.",
        "Terven, J., C\u00f3rdova-Esparza, D. M., & Romero-Gonz\u00e1lez, J. A. (2023). A comprehensive review of YOLO architectures in computer vision: From YOLOv1 to YOLOv8 and YOLO-NAS. Machine Learning and Knowledge Extraction, 5(4), 1680-1716. https://doi.org/10.3390/make5040083",
        "Thiagarajan, S., Semmel, D. S., & Semmel, M. I. (1974). Instructional development for training teachers of exceptional children. National Center for Improvement of Educational Systems.",
    ]

    for ref in references:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
        run = para.add_run(ref)
        set_run_font(run, size=12)

    # ── Save ──
    output_path = r"d:\project-amsp\Proposal_Skripsi_LMS_AI_Proctoring.docx"
    doc.save(output_path)
    print(f"\n[OK] File berhasil dibuat: {output_path}")
    print("Silakan buka file tersebut di Microsoft Word.")


if __name__ == "__main__":
    build_proposal()
